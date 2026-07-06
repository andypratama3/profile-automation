#!/usr/bin/env python3
"""
Generate CV ATS dari profile.json — Word & PDF.

Usage:
    python generate_cv.py --data data/profile.json --output cv_saya.docx
    python generate_cv.py --data data/profile.json --output cv_saya.pdf --format pdf

Requires: pip install python-docx
"""

import json, argparse, os

def load_json(path):
    with open(path) as f:
        return json.load(f)

def generate(profile):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # HEADER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(profile.get("name", "Nama Kamu"))
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    kontak = [profile.get("email"), profile.get("phone"), profile.get("linkedin_url")]
    kontak = [k for k in kontak if k]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(" | ".join(kontak))
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ABOUT
    if profile.get("about"):
        p = doc.add_paragraph()
        r = p.add_run(profile["about"])
        r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        border = p.paragraph_format._element.get_or_add_pPr()
        pBdr = border.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "1", qn("w:color"): "1A56DB"})
        pBdr.append(bottom); border.append(pBdr)
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    def add_job(title, company, dates, loc, bullets):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title); r.bold = True; r.font.size = Pt(12)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(f"{company} | {loc} | {dates}")
        r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for b in bullets:
            bp = doc.add_paragraph(b, style='List Bullet')
            bp.paragraph_format.space_after = Pt(1)
            for r in bp.runs: r.font.size = Pt(10)

    # EXPERIENCE
    if profile.get("work_experience"):
        heading("Pengalaman")
        for job in profile["work_experience"]:
            add_job(
                job.get("position", ""), job.get("company", ""),
                f"{job.get('start_date', '')} - {job.get('end_date', '')}",
                job.get("location", ""), job.get("bullets", []))

    # EDUCATION
    if profile.get("education"):
        heading("Pendidikan")
        for edu in profile["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(edu.get("degree", edu.get("major", "")))
            r.bold = True; r.font.size = Pt(11)
            p2 = doc.add_paragraph()
            info = f"{edu.get('institution', '')} | {edu.get('start_year', '')} - {edu.get('end_year', '')}"
            if edu.get("gpa"): info += f" | IPK: {edu['gpa']}"
            r2 = p2.add_run(info); r2.font.size = Pt(10.5)
            r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # CERTIFICATIONS
    if profile.get("certifications"):
        heading("Sertifikasi")
        for c in profile["certifications"]:
            if isinstance(c, str): text = c
            else: text = f"{c.get('name','')} - {c.get('publisher','')} ({c.get('year','')})"
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs: r.font.size = Pt(10)

    # SKILLS
    skills = profile.get("hard_skills") or profile.get("skills") or []
    if skills:
        heading("Keahlian")
        p = doc.add_paragraph(", ".join(skills))
        for r in p.runs: r.font.size = Pt(10)

    return doc

def main():
    parser = argparse.ArgumentParser(description="Generate CV ATS dari profile.json")
    parser.add_argument("--data", default="data/profile.json")
    parser.add_argument("--output", default="cv_output.docx")
    args = parser.parse_args()

    profile = load_json(args.data)
    doc = generate(profile)
    out = args.output if args.output.endswith(".docx") else args.output + ".docx"
    doc.save(out)
    print(f"✅ CV tersimpan: {out}")

if __name__ == "__main__":
    main()
